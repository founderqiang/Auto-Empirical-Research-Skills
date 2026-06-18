#!/usr/bin/env python3
"""
Markdown to DOCX Converter
Converts markdown files to Word documents with proper formatting.

Usage:
    python convert_md_to_docx.py input.md output.docx [options]
"""

import sys
import os
import argparse
import subprocess
import re
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from copy import deepcopy

def check_pandoc_installed():
    """Check if pandoc is installed on the system."""
    try:
        result = subprocess.run(['pandoc', '--version'],
                              capture_output=True,
                              text=True)
        return result.returncode == 0
    except FileNotFoundError:
        return False

def preprocess_markdown(content):
    """
    预处理Markdown内容，将单个换行转换为双换行，使其成为独立段落。
    保留已有的双换行（空行）不变。

    Args:
        content (str): 原始Markdown内容

    Returns:
        str: 处理后的Markdown内容
    """
    lines = content.split('\n')
    processed_lines = []
    i = 0

    while i < len(lines):
        current_line = lines[i].rstrip()

        # 检查下一行是否存在
        if i + 1 < len(lines):
            next_line = lines[i + 1].rstrip()

            # 如果当前行不是空行，下一行也不是空行
            # 并且当前行不是列表项、标题、代码块等特殊格式
            if (current_line and next_line and
                not current_line.startswith('#') and  # 不是标题
                not current_line.startswith('-') and  # 不是列表
                not current_line.startswith('*') and  # 不是列表
                not current_line.startswith('+') and  # 不是列表
                not re.match(r'^\d+\.', current_line) and  # 不是有序列表
                not current_line.startswith('```') and  # 不是代码块
                not current_line.startswith('|')):  # 不是表格

                # 添加当前行和一个空行
                processed_lines.append(current_line)
                processed_lines.append('')  # 添加空行使其成为独立段落
            else:
                # 保持原样
                processed_lines.append(current_line)
        else:
            # 最后一行，保持原样
            processed_lines.append(current_line)

        i += 1

    return '\n'.join(processed_lines)

def fix_heading_fonts(docx_file):
    """
    修复Word文档中所有标题的字体，确保使用仿宋，并设置1.5倍行距和段后间距

    Args:
        docx_file (str): Word文档路径

    Returns:
        bool: 成功返回True，失败返回False
    """
    try:
        # 打开文档
        doc = Document(docx_file)

        # 遍历所有段落
        for para in doc.paragraphs:
            # 为所有段落设置1.5倍行距
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

            # 检查是否是标题段落
            if para.style.name.startswith('Heading') or para.style.name in ['Title', 'Subtitle']:
                # 设置段后间距为1行（12磅）
                para.paragraph_format.space_after = Pt(12)

                # 为该段落的每个run设置字体
                for run in para.runs:
                    run.font.name = '仿宋'
                    run.font.italic = False  # 确保不是斜体
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色

                    # 设置所有字体属性
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                    r.rPr.rFonts.set(qn('w:ascii'), '仿宋')
                    r.rPr.rFonts.set(qn('w:hAnsi'), '仿宋')
                    r.rPr.rFonts.set(qn('w:cs'), '仿宋')

        # 保存文档
        doc.save(docx_file)
        return True

    except Exception as e:
        print(f"Warning: Could not fix heading fonts and line spacing: {str(e)}")
        return False

def fix_table_borders(docx_file):
    """
    为Word文档中所有表格添加黑色单线边框（外框+内框）。

    Args:
        docx_file (str): Word文档路径

    Returns:
        int: 修复的表格数量
    """
    try:
        doc = Document(docx_file)
        count = 0
        for table in doc.tables:
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                from docx.oxml import OxmlElement
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)

            # Remove existing tblBorders if any
            for existing in tblPr.findall(qn('w:tblBorders')):
                tblPr.remove(existing)

            # Create tblBorders element
            from docx.oxml import OxmlElement
            borders = OxmlElement('w:tblBorders')
            for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                element = OxmlElement(f'w:{edge}')
                element.set(qn('w:val'), 'single')
                element.set(qn('w:sz'), '4')
                element.set(qn('w:space'), '0')
                element.set(qn('w:color'), '000000')
                borders.append(element)

            # Insert tblBorders after tblW (proper schema order)
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is not None:
                tblW.addnext(borders)
            else:
                tblPr.append(borders)

            count += 1

        doc.save(docx_file)
        return count

    except Exception as e:
        print(f"Warning: Could not fix table borders: {str(e)}")
        return 0


def is_chinese_char(char):
    """检查字符是否为中文字符"""
    return '\u4e00' <= char <= '\u9fff'

def convert_citations_to_superscript(docx_file):
    """
    将文档中的引用数字（如句末的1、2、3等）转换为上标格式。

    识别规则（更严格）：
    - 1-2位数字，紧跟在中文字符或英文字母后面
    - 数字后面紧跟句末标点（。；）
    - 排除年份（如2030年）、百分比（60%）、小数（2.9）、序号（1.1）等
    - 排除表格编号（表 1）、列表序号等
    - 排除大于60的数字

    Args:
        docx_file (str): Word文档路径

    Returns:
        int: 处理的段落数量
    """
    try:
        doc = Document(docx_file)

        # 不应该转换为上标的单位和后缀
        excluded_suffixes = {'年', '月', '日', '万', '亿', '%', '个', '位', '次',
                            '条', '项', '种', '类', '倍', '层', '步', '点', '章',
                            '节', '篇', '页', '行', '列', '组', '期', '号', '版'}

        def process_paragraph(para):
            """处理单个段落，将引用数字转换为上标"""
            full_text = para.text
            if not full_text or not para.runs:
                return False

            # 跳过标题、表格标题行等
            style_name = para.style.name if para.style else ''
            if style_name.startswith('Heading') or style_name in ['Title', 'Subtitle']:
                return False

            # 匹配：数字后跟句末标点（。；）
            pattern = r'(\d{1,2})([。；])'

            matches = list(re.finditer(pattern, full_text))
            if not matches:
                return False

            citation_positions = []
            for match in matches:
                num = match.group(1)
                pos = match.start()
                end_pos = match.end() - 1  # 不包括标点

                # 检查数字前面的字符
                if pos > 0:
                    char_before = full_text[pos - 1]
                    # 必须紧跟在中文字符或英文字母后面（不能是数字、空格、标点）
                    if not (is_chinese_char(char_before) or char_before.isalpha()):
                        continue
                    # 排除小数点前的数字（如 2.9）
                    if char_before == '.':
                        continue
                else:
                    # 数字在段落开头，不太可能是引用
                    continue

                # 检查是否是更大数字的一部分（如年份2030中的30）
                if pos >= 2:
                    two_before = full_text[pos - 2:pos]
                    if two_before.isdigit() or (two_before[0].isdigit() and two_before[1] == '.'):
                        continue

                # 排除过大的数字
                if int(num) > 60:
                    continue

                citation_positions.append((pos, pos + len(num), num))

            if not citation_positions:
                return False

            # 重建段落，将引用数字设为上标
            original_text = full_text

            # 清空现有runs
            for run in list(para.runs):
                para._p.remove(run._r)

            # 重建段落
            last_pos = 0
            for start, end, num in citation_positions:
                if start > last_pos:
                    para.add_run(original_text[last_pos:start])
                run = para.add_run(num)
                run.font.superscript = True
                last_pos = end

            if last_pos < len(original_text):
                para.add_run(original_text[last_pos:])

            return True

        # 处理所有段落
        modified_count = 0
        for para in doc.paragraphs:
            if process_paragraph(para):
                modified_count += 1

        # 处理表格中的段落
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if process_paragraph(para):
                            modified_count += 1

        doc.save(docx_file)
        return modified_count

    except Exception as e:
        print(f"Warning: Could not convert citations to superscript: {str(e)}")
        return 0

def convert_md_to_docx(input_file, output_file, reference_doc=None,
                       toc=False, standalone=True, metadata=None,
                       use_chinese_template=True, superscript_citations=True):
    """
    Convert a markdown file to DOCX using pandoc.

    Args:
        input_file (str): Path to input markdown file
        output_file (str): Path to output DOCX file
        reference_doc (str, optional): Path to reference DOCX for styling
        toc (bool): Generate table of contents
        standalone (bool): Create standalone document
        metadata (dict, optional): Document metadata (title, author, date)
        use_chinese_template (bool): Use default Chinese template with FangSong font
        superscript_citations (bool): Convert citation numbers to superscript (default: True)

    Returns:
        bool: True if conversion successful, False otherwise
    """
    # Check if pandoc is installed
    if not check_pandoc_installed():
        print("Error: pandoc is not installed.")
        print("Please install pandoc:")
        print("  macOS:   brew install pandoc")
        print("  Linux:   sudo apt-get install pandoc")
        print("  Windows: choco install pandoc")
        return False

    # Check if input file exists
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found.")
        return False

    # 读取并预处理Markdown文件
    with open(input_file, 'r', encoding='utf-8') as f:
        original_content = f.read()

    # 预处理：将单个换行转换为段落分隔
    processed_content = preprocess_markdown(original_content)

    # 创建临时文件保存处理后的内容
    temp_fd, temp_file = tempfile.mkstemp(suffix='.md', text=True)
    try:
        with os.fdopen(temp_fd, 'w', encoding='utf-8') as f:
            f.write(processed_content)

        # Use default Chinese template if no reference doc specified
        if reference_doc is None and use_chinese_template:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            default_template = os.path.join(script_dir, 'chinese_template.docx')
            if os.path.exists(default_template):
                reference_doc = default_template
                print(f"使用中文模板（仿宋字体）")

        # Build pandoc command (使用临时文件而不是原始文件)
        cmd = ['pandoc', temp_file, '-o', output_file]

        # Add options
        if standalone:
            cmd.append('--standalone')

        if toc:
            cmd.append('--toc')

        if reference_doc and os.path.exists(reference_doc):
            cmd.extend(['--reference-doc', reference_doc])

        # Add metadata if provided
        if metadata:
            if 'title' in metadata:
                cmd.extend(['--metadata', f'title={metadata["title"]}'])
            if 'author' in metadata:
                cmd.extend(['--metadata', f'author={metadata["author"]}'])
            if 'date' in metadata:
                cmd.extend(['--metadata', f'date={metadata["date"]}'])

        # Run conversion
        try:
            print(f"Converting '{input_file}' to '{output_file}'...")
            result = subprocess.run(cmd, capture_output=True, text=True)

            if result.returncode == 0:
                print(f"✓ Successfully converted to '{output_file}'")

                # 后处理：修复标题字体、行距和段后间距
                if use_chinese_template:
                    print(f"正在修复标题字体、行距和段后间距...")
                    if fix_heading_fonts(output_file):
                        print(f"✓ 标题字体已修复为仿宋，行距1.5倍，标题段后1行")

                # 后处理：为表格添加边框
                print(f"正在为表格添加边框...")
                table_count = fix_table_borders(output_file)
                if table_count > 0:
                    print(f"✓ 已为 {table_count} 个表格添加边框")

                # 后处理：将引用数字转换为上标
                if superscript_citations:
                    print(f"正在将引用数字转换为上标...")
                    count = convert_citations_to_superscript(output_file)
                    if count > 0:
                        print(f"✓ 已将 {count} 个段落中的引用数字转换为上标")

                return True
            else:
                print(f"Error during conversion:")
                print(result.stderr)
                return False

        except Exception as e:
            print(f"Error: {str(e)}")
            return False

    finally:
        # 清理临时文件
        if os.path.exists(temp_file):
            os.unlink(temp_file)

def batch_convert(input_dir, output_dir, pattern='*.md', **kwargs):
    """
    Batch convert all markdown files in a directory.

    Args:
        input_dir (str): Input directory path
        output_dir (str): Output directory path
        pattern (str): File pattern to match (default: '*.md')
        **kwargs: Additional arguments for convert_md_to_docx
    """
    input_path = Path(input_dir)
    output_path = Path(output_dir)

    # Create output directory if it doesn't exist
    output_path.mkdir(parents=True, exist_ok=True)

    # Find all markdown files
    md_files = list(input_path.glob(pattern))

    if not md_files:
        print(f"No markdown files found in '{input_dir}'")
        return

    print(f"Found {len(md_files)} markdown file(s)")

    # Convert each file
    success_count = 0
    for md_file in md_files:
        docx_file = output_path / f"{md_file.stem}.docx"
        if convert_md_to_docx(str(md_file), str(docx_file), **kwargs):
            success_count += 1

    print(f"\n{success_count}/{len(md_files)} files converted successfully")

def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown files to Word documents (DOCX)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Convert a single file (default: no table of contents)
  python convert_md_to_docx.py README.md README.docx

  # Convert with custom styling template
  python convert_md_to_docx.py input.md output.docx --reference-doc template.docx

  # Convert with table of contents (optional)
  python convert_md_to_docx.py report.md report.docx --toc

  # Batch convert all .md files in a directory
  python convert_md_to_docx.py --batch input_dir/ output_dir/
        """
    )

    parser.add_argument('input', help='Input markdown file or directory (for batch mode)')
    parser.add_argument('output', nargs='?', help='Output DOCX file or directory (for batch mode)')
    parser.add_argument('--reference-doc', help='Reference DOCX file for styling')
    parser.add_argument('--toc', action='store_true', help='Generate table of contents (default: no TOC)')
    parser.add_argument('--no-standalone', action='store_true', help='Do not create standalone document')
    parser.add_argument('--no-chinese-template', action='store_true', help='Do not use default Chinese template (FangSong font)')
    parser.add_argument('--no-superscript-citations', action='store_true', help='Do not convert citation numbers to superscript')
    parser.add_argument('--title', help='Document title')
    parser.add_argument('--author', help='Document author')
    parser.add_argument('--date', help='Document date')
    parser.add_argument('--batch', action='store_true', help='Batch convert all .md files in directory')
    parser.add_argument('--pattern', default='*.md', help='File pattern for batch mode (default: *.md)')

    args = parser.parse_args()

    # Prepare metadata
    metadata = {}
    if args.title:
        metadata['title'] = args.title
    if args.author:
        metadata['author'] = args.author
    if args.date:
        metadata['date'] = args.date

    # Batch mode
    if args.batch:
        if not args.output:
            print("Error: Output directory required for batch mode")
            sys.exit(1)

        batch_convert(
            args.input,
            args.output,
            pattern=args.pattern,
            reference_doc=args.reference_doc,
            toc=args.toc,
            standalone=not args.no_standalone,
            metadata=metadata if metadata else None,
            use_chinese_template=not args.no_chinese_template,
            superscript_citations=not args.no_superscript_citations
        )
    else:
        # Single file mode
        if not args.output:
            print("Error: Output file required")
            parser.print_help()
            sys.exit(1)

        success = convert_md_to_docx(
            args.input,
            args.output,
            reference_doc=args.reference_doc,
            toc=args.toc,
            standalone=not args.no_standalone,
            metadata=metadata if metadata else None,
            use_chinese_template=not args.no_chinese_template,
            superscript_citations=not args.no_superscript_citations
        )

        sys.exit(0 if success else 1)

if __name__ == '__main__':
    main()
