#!/usr/bin/env python3
"""Fix table formatting in docx files - improved version."""

import sys
from docx import Document
from docx.shared import Pt, Cm, Twips, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_cell_border(cell, **kwargs):
    """Set cell border."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove existing borders
    for tcBorders in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(tcBorders)

    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            element.set(qn('w:space'), str(edge_data.get('space', 0)))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_width(cell, width_twips):
    """Set cell width explicitly."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove existing width
    for tcW in tcPr.findall(qn('w:tcW')):
        tcPr.remove(tcW)

    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_twips)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_table_layout_fixed(table):
    """Set table layout to fixed."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove existing layout
    for tblLayout in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(tblLayout)

    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

def set_table_width(table, width_twips):
    """Set table width."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove existing width
    for tblW in tblPr.findall(qn('w:tblW')):
        tblPr.remove(tblW)

    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(width_twips)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def set_table_grid(table, col_widths_twips):
    """Set table grid column widths."""
    tbl = table._tbl

    # Remove existing grid
    for tblGrid in tbl.findall(qn('w:tblGrid')):
        tbl.remove(tblGrid)

    tblGrid = OxmlElement('w:tblGrid')
    for width in col_widths_twips:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(width)))
        tblGrid.append(gridCol)

    # Insert grid after tblPr
    tblPr = tbl.tblPr
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)

def fix_tables(doc_path, output_path=None):
    """Fix all tables in the document."""
    doc = Document(doc_path)

    # Page width minus margins (A4: 21cm - 2.54cm*2 ≈ 16cm = 9072 twips)
    # 1 cm = 567 twips, 1 inch = 1440 twips
    page_width_twips = 9072

    for table_idx, table in enumerate(doc.tables):
        num_cols = len(table.columns)
        print(f"Processing table {table_idx + 1} with {num_cols} columns")

        # Calculate column widths based on content
        if num_cols == 2:
            # For 2-column tables like "符号 | 含义"
            col_widths_twips = [2500, 6572]  # ~4.4cm, ~11.6cm
        elif num_cols == 3:
            # For 3-column tables like "理论预测 | 中国平台实践 | 验证程度"
            col_widths_twips = [2200, 5372, 1500]  # ~3.9cm, ~9.5cm, ~2.6cm
        else:
            # Equal distribution
            col_widths_twips = [page_width_twips // num_cols] * num_cols

        total_width = sum(col_widths_twips)

        # Set table properties
        set_table_layout_fixed(table)
        set_table_width(table, total_width)
        set_table_grid(table, col_widths_twips)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Process each row and cell
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # Set cell width
                set_cell_width(cell, col_widths_twips[col_idx])

                # Set cell borders
                border_style = {'val': 'single', 'sz': 4, 'color': '000000'}
                set_cell_border(cell,
                    top=border_style,
                    bottom=border_style,
                    left=border_style,
                    right=border_style
                )

                # Set cell margins/padding
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                # Remove existing margins
                for tcMar in tcPr.findall(qn('w:tcMar')):
                    tcPr.remove(tcMar)

                tcMar = OxmlElement('w:tcMar')
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), '80')  # 80 twips padding
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                tcPr.append(tcMar)

                # Set vertical alignment to center
                for vAlign in tcPr.findall(qn('w:vAlign')):
                    tcPr.remove(vAlign)
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)

                # Process paragraphs in cell
                for para in cell.paragraphs:
                    # Remove first line indent for table cells
                    para.paragraph_format.first_line_indent = Pt(0)
                    # Set line spacing
                    para.paragraph_format.line_spacing = 1.15
                    para.paragraph_format.space_before = Pt(3)
                    para.paragraph_format.space_after = Pt(3)
                    # Left align for better readability
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Bold header row
                    if row_idx == 0:
                        for run in para.runs:
                            run.bold = True

    # Save the document
    output = output_path or doc_path
    doc.save(output)
    print(f"✓ 表格格式已修复：{output}")

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python fix_tables.py input.docx [output.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    fix_tables(input_file, output_file)
