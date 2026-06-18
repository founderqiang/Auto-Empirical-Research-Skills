#!/usr/bin/env python3
"""
创建中文Word模板
设置默认中文字体为仿宋，英文字体为Times New Roman
默认行距：1.5倍
"""

import sys
import os

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.enum.text import WD_LINE_SPACING
except ImportError:
    print("错误：需要安装 python-docx 库")
    print("请运行：pip install python-docx")
    sys.exit(1)

def create_chinese_template(output_file="chinese_template.docx"):
    """
    创建一个中文Word模板文档
    - 中文字体：仿宋（正文和所有标题，包括副标题）
    - 英文字体：Times New Roman
    - 字体颜色：黑色
    - 首行缩进：两个字符（24磅，精确对齐）
    - 标题样式：不使用斜体
    - 行距：1.5倍
    """

    doc = Document()

    # 设置文档的默认样式
    styles = doc.styles

    # 正文样式
    style_normal = styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(12)
    font_normal.color.rgb = RGBColor(0, 0, 0)  # 黑色

    # 设置首行缩进两个字符（精确计算：2字符 × 12磅字号）
    paragraph_format = style_normal.paragraph_format
    paragraph_format.first_line_indent = Pt(24)  # 2个字符 = 2 × 12磅 = 24磅

    # 设置1.5倍行距
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # 设置中文字体为仿宋
    element = style_normal.element
    element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    # 标题1样式 - 仿宋，不斜体，1.5倍行距，段后1行
    style_h1 = styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = '仿宋'  # 直接设置为仿宋
    font_h1.size = Pt(22)
    font_h1.bold = True
    font_h1.italic = False  # 明确不使用斜体
    font_h1.color.rgb = RGBColor(0, 0, 0)  # 黑色
    # 彻底设置所有字体属性
    element_h1 = style_h1.element
    element_h1.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    element_h1.rPr.rFonts.set(qn('w:ascii'), '仿宋')
    element_h1.rPr.rFonts.set(qn('w:hAnsi'), '仿宋')
    element_h1.rPr.rFonts.set(qn('w:cs'), '仿宋')
    # 清除段落格式中的首行缩进，设置1.5倍行距，段后1行
    style_h1.paragraph_format.first_line_indent = Pt(0)
    style_h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style_h1.paragraph_format.space_after = Pt(12)  # 段后1行间距（12磅）

    # 标题2样式 - 仿宋，不斜体，1.5倍行距，段后1行
    style_h2 = styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = '仿宋'  # 直接设置为仿宋
    font_h2.size = Pt(18)
    font_h2.bold = True
    font_h2.italic = False  # 明确不使用斜体
    font_h2.color.rgb = RGBColor(0, 0, 0)  # 黑色
    element_h2 = style_h2.element
    element_h2.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    element_h2.rPr.rFonts.set(qn('w:ascii'), '仿宋')
    element_h2.rPr.rFonts.set(qn('w:hAnsi'), '仿宋')
    element_h2.rPr.rFonts.set(qn('w:cs'), '仿宋')
    style_h2.paragraph_format.first_line_indent = Pt(0)
    style_h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style_h2.paragraph_format.space_after = Pt(12)  # 段后1行间距

    # 标题3样式 - 仿宋，不斜体，1.5倍行距，段后1行
    style_h3 = styles['Heading 3']
    font_h3 = style_h3.font
    font_h3.name = '仿宋'  # 直接设置为仿宋
    font_h3.size = Pt(16)
    font_h3.bold = True
    font_h3.italic = False  # 明确不使用斜体
    font_h3.color.rgb = RGBColor(0, 0, 0)  # 黑色
    element_h3 = style_h3.element
    element_h3.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    element_h3.rPr.rFonts.set(qn('w:ascii'), '仿宋')
    element_h3.rPr.rFonts.set(qn('w:hAnsi'), '仿宋')
    element_h3.rPr.rFonts.set(qn('w:cs'), '仿宋')
    style_h3.paragraph_format.first_line_indent = Pt(0)
    style_h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style_h3.paragraph_format.space_after = Pt(12)  # 段后1行间距

    # 标题4样式 - 仿宋，不斜体，1.5倍行距，段后1行
    style_h4 = styles['Heading 4']
    font_h4 = style_h4.font
    font_h4.name = '仿宋'  # 直接设置为仿宋
    font_h4.size = Pt(14)
    font_h4.bold = True
    font_h4.italic = False  # 明确不使用斜体
    font_h4.color.rgb = RGBColor(0, 0, 0)  # 黑色
    element_h4 = style_h4.element
    element_h4.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    element_h4.rPr.rFonts.set(qn('w:ascii'), '仿宋')
    element_h4.rPr.rFonts.set(qn('w:hAnsi'), '仿宋')
    element_h4.rPr.rFonts.set(qn('w:cs'), '仿宋')
    style_h4.paragraph_format.first_line_indent = Pt(0)
    style_h4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style_h4.paragraph_format.space_after = Pt(12)  # 段后1行间距

    # 标题5样式 - 仿宋，不斜体，1.5倍行距，段后1行
    style_h5 = styles['Heading 5']
    font_h5 = style_h5.font
    font_h5.name = '仿宋'
    font_h5.size = Pt(13)
    font_h5.bold = True
    font_h5.italic = False
    font_h5.color.rgb = RGBColor(0, 0, 0)
    element_h5 = style_h5.element
    element_h5.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    element_h5.rPr.rFonts.set(qn('w:ascii'), '仿宋')
    element_h5.rPr.rFonts.set(qn('w:hAnsi'), '仿宋')
    element_h5.rPr.rFonts.set(qn('w:cs'), '仿宋')
    style_h5.paragraph_format.first_line_indent = Pt(0)
    style_h5.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style_h5.paragraph_format.space_after = Pt(12)  # 段后1行间距

    # 标题6样式 - 仿宋，不斜体，1.5倍行距，段后1行
    style_h6 = styles['Heading 6']
    font_h6 = style_h6.font
    font_h6.name = '仿宋'
    font_h6.size = Pt(12)
    font_h6.bold = True
    font_h6.italic = False
    font_h6.color.rgb = RGBColor(0, 0, 0)
    element_h6 = style_h6.element
    element_h6.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    element_h6.rPr.rFonts.set(qn('w:ascii'), '仿宋')
    element_h6.rPr.rFonts.set(qn('w:hAnsi'), '仿宋')
    element_h6.rPr.rFonts.set(qn('w:cs'), '仿宋')
    style_h6.paragraph_format.first_line_indent = Pt(0)
    style_h6.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style_h6.paragraph_format.space_after = Pt(12)  # 段后1行间距

    # Title样式（文档标题）- 仿宋，不斜体，1.5倍行距，段后1行
    if 'Title' in styles:
        style_title = styles['Title']
        font_title = style_title.font
        font_title.name = '仿宋'
        font_title.size = Pt(26)
        font_title.bold = True
        font_title.italic = False
        font_title.color.rgb = RGBColor(0, 0, 0)
        element_title = style_title.element
        element_title.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        element_title.rPr.rFonts.set(qn('w:ascii'), '仿宋')
        element_title.rPr.rFonts.set(qn('w:hAnsi'), '仿宋')
        element_title.rPr.rFonts.set(qn('w:cs'), '仿宋')
        style_title.paragraph_format.first_line_indent = Pt(0)
        style_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style_title.paragraph_format.space_after = Pt(12)  # 段后1行间距

    # Subtitle样式（副标题）- 仿宋，不斜体，1.5倍行距，段后1行
    if 'Subtitle' in styles:
        style_subtitle = styles['Subtitle']
        font_subtitle = style_subtitle.font
        font_subtitle.name = '仿宋'
        font_subtitle.size = Pt(14)
        font_subtitle.bold = False
        font_subtitle.italic = False  # 明确不使用斜体
        font_subtitle.color.rgb = RGBColor(0, 0, 0)
        element_subtitle = style_subtitle.element
        element_subtitle.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        element_subtitle.rPr.rFonts.set(qn('w:ascii'), '仿宋')
        element_subtitle.rPr.rFonts.set(qn('w:hAnsi'), '仿宋')
        element_subtitle.rPr.rFonts.set(qn('w:cs'), '仿宋')
        style_subtitle.paragraph_format.first_line_indent = Pt(0)
        style_subtitle.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style_subtitle.paragraph_format.space_after = Pt(12)  # 段后1行间距

    # 添加示例内容（可选）
    doc.add_heading('中文文档模板', 0)

    p = doc.add_paragraph('本模板默认使用以下字体设置：')
    p = doc.add_paragraph('中文：仿宋（正文和所有标题）')
    p = doc.add_paragraph('英文：Times New Roman')
    p = doc.add_paragraph('字号：12磅（正文）')
    p = doc.add_paragraph('字体颜色：黑色')
    p = doc.add_paragraph('首行缩进：两个字符（24磅，精确对齐）')
    p = doc.add_paragraph('标题样式：不使用斜体')

    doc.add_heading('使用说明', 1)
    p = doc.add_paragraph('此模板文件用于 Markdown 转 Word 转换时的样式参考。')
    p = doc.add_paragraph('转换时使用 --reference-doc 参数指定此模板即可。')
    p = doc.add_paragraph('所有标题（H1-H6）和副标题均使用仿宋字体，不使用斜体。')

    # 保存文档
    doc.save(output_file)
    print(f"✓ 中文模板已创建：{output_file}")
    print(f"  - 中文字体：仿宋（正文和所有标题，包括副标题）")
    print(f"  - 英文字体：Times New Roman")
    print(f"  - 正文字号：12磅")
    print(f"  - 字体颜色：黑色")
    print(f"  - 首行缩进：两个字符（24磅，精确对齐）")
    print(f"  - 标题样式：不使用斜体")
    print(f"  - 行距：1.5倍")
    print(f"  - 标题段后间距：1行")

    return True

def main():
    """主函数"""
    import argparse

    parser = argparse.ArgumentParser(description='创建中文Word模板')
    parser.add_argument('-o', '--output',
                       default='chinese_template.docx',
                       help='输出文件名（默认：chinese_template.docx）')

    args = parser.parse_args()

    try:
        create_chinese_template(args.output)
    except Exception as e:
        print(f"错误：{str(e)}")
        sys.exit(1)

if __name__ == '__main__':
    main()
